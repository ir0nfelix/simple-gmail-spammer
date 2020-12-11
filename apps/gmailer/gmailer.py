import glob
import pathlib
from operator import itemgetter
from time import sleep

import xlrd
from docx import Document
from envelopes import Envelope, GMailSMTP
from imap_tools import MailBox, OR, A
from xlutils.copy import copy


class GMailer:
    def __init__(self, *args, **kwargs):
        self.email_address = kwargs.get('email_address', 'my_default_gmail@gmail.com')
        self.password = kwargs.get('password', 'my_default_password')
        self.recepient_xls_path = kwargs.get('recepient_xls_path', 'files/my_list_of_recepients.xls')
        self.attach_file_path = kwargs.get('attach_file_path', 'files/my_attachment.pdf')

    def _get_messages_by_email(self, folder: str, recepient_email: str) -> list:
        with MailBox('smtp.gmail.com').login(self.email_address, self.password, f'{folder}') as mailbox:
            message_list = [(msg.date, msg.text or msg.html) for msg in mailbox.fetch(OR(A(from_=f'{recepient_email}'),
                                                                                         A(to=f'{recepient_email}')))]

        return message_list

    def _get_recepient_mail_chain(self, recepient_email: str) -> list:
        inbox_msg_list = self._get_messages_by_email('INBOX', recepient_email)
        sent_msg_list = self._get_messages_by_email('[Gmail]/Sent Mail', recepient_email)
        return [i[1] for i in sorted(inbox_msg_list + sent_msg_list, key=itemgetter(0))]

    def _get_chain_file_name(self, recepient_email: str) -> str:
        inbox_msg_list = self._get_messages_by_email('INBOX', recepient_email)
        sent_msg_list = self._get_messages_by_email('[Gmail]/Sent Mail', recepient_email)
        date_str = max([i[0] for i in inbox_msg_list + sent_msg_list]).strftime('%d-%m-%Y_%H:%M')
        return f'files/{recepient_email}_{date_str}.docx'

    def _get_recepient_from_xls_file(self) -> list:
        xls_workbook = xlrd.xlrd.open_workbook(self.recepient_xls_path)
        recepient_list = []
        for index in range(0, xls_workbook.nsheets):
            xls_sheet = xls_workbook.sheet_by_index(index)
            for row in range(1, xls_sheet.nrows):
                recepient_list.append((xls_sheet.cell(row, 0).value,
                                       xls_sheet.cell(row, 4).value,
                                       bool(xls_sheet.cell(row, 14).value),
                                       bool(xls_sheet.cell(row, 15).value),
                                       index,
                                       row))
        recepient_list_is_needed = list(filter(lambda x: x[2] is True, recepient_list))
        return recepient_list_is_needed

    def _create_chain_file(self, recepient_email: str):
        message_chain = self._get_recepient_mail_chain(recepient_email)
        if not message_chain or len(message_chain) < 2:
            return

        document = Document()
        for msg in message_chain:
            document.add_paragraph(msg)
            document.add_paragraph(60*'*')

        file_name = self._get_chain_file_name(recepient_email)
        existing_file = glob.glob('files/' + f'{recepient_email}*.docx')

        if file_name in existing_file:
            return

        if existing_file:
            [path.unlink() for path in [pathlib.Path(file) for file in existing_file]]

        document.save(f'{file_name}')

    def update_chains(self):
        recepient_list_is_needed = self._get_recepient_from_xls_file()
        [self._create_chain_file(recepient[1]) for recepient in recepient_list_is_needed if recepient[3]]

    def send_mails(self):
        recepient_list_not_send = list(filter(lambda x: x[3] is False, self._get_recepient_from_xls_file()))

        if not recepient_list_not_send:
            return

        for recepient in recepient_list_not_send:
            text_body = f'''
                Dear client {recepient[0]},
                ````My text````
                
                Best regards, My Name.
                '''
            try:
                envelope = Envelope(
                    from_addr='My Name',
                    to_addr=recepient[1],
                    subject='My Subject',
                    text_body=text_body)
                envelope.add_attachment(self.attach_file_path)
                gmail = GMailSMTP(login=self.email_address, password=self.password)
                gmail.send(envelope)

                rb = xlrd.open_workbook(self.recepient_xls_path)
                wb = copy(rb)
                s = wb.get_sheet(recepient[4])
                s.write(recepient[5], 15, 1)
                wb.save(self.recepient_xls_path)
                sleep(5)

            except Exception:
                rb = xlrd.open_workbook(self.recepient_xls_path)
                wb = copy(rb)
                s = wb.get_sheet(recepient[4])
                s.write(recepient[5], 15, 0)
                wb.save(self.recepient_xls_path)
                continue
